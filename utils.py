import os
import gc
import re
from io import BytesIO
from zipfile import ZipFile

import numpy as np
import streamlit as st
from PIL import Image
from docx import Document
import pdfplumber
import fitz

def format_datetime(iso_str):
    try:
        dt = datetime.fromisoformat(iso_str)
        return dt.strftime("%B %d, %Y at %I:%M %p %Z") or dt.strftime("%B %d, %Y at %I:%M %p UTC")
    except Exception:
        return iso_str

from datetime import datetime, timedelta, timezone
import re

def decode_pdf_date(date_str):
    if not date_str:
        return None

    try:
        # Remove 'D:' prefix if present
        if date_str.startswith("D:"):
            date_str = date_str[2:]

        # Match basic parts
        pattern = r"(?P<year>\d{4})(?P<month>\d{2})(?P<day>\d{2})" \
                  r"(?P<hour>\d{2})(?P<minute>\d{2})(?P<second>\d{2})?" \
                  r"(?P<tz_sign>[+-Z])?(?P<tz_hour>\d{2})?'?(?P<tz_minute>\d{2})?'?"

        match = re.match(pattern, date_str)
        if not match:
            return None

        parts = match.groupdict()
        dt = datetime(
            int(parts["year"]),
            int(parts["month"]),
            int(parts["day"]),
            int(parts["hour"]),
            int(parts["minute"]),
            int(parts.get("second") or 0)
        )

        if parts["tz_sign"] in ["+", "-"]:
            tz_offset = timedelta(
                hours=int(parts["tz_hour"] or 0),
                minutes=int(parts["tz_minute"] or 0)
            )
            if parts["tz_sign"] == "-":
                tz_offset = -tz_offset
            dt = dt.replace(tzinfo=timezone(tz_offset))
        elif parts["tz_sign"] == "Z":
            dt = dt.replace(tzinfo=timezone.utc)

        return format_datetime(dt.isoformat())
    except Exception:
        return None


def run_easyocr(img_np):
    import easyocr
    try:
        reader = easyocr.Reader(['en'], gpu=False, verbose=False)
        result = reader.readtext(img_np)
        del reader
        gc.collect()
        return result
    except Exception as e:
        st.warning(f"OCR reader failed: {e}")
        return []


def pdf_extract(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
    except Exception as e:
        st.error(f"PDF text extraction failed: {e}")
        return ""


def ocr_pdf_extract(file_path, max_pages=3):
    text = ""
    try:
        doc = fitz.open(file_path)
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            try:
                pix = page.get_pixmap(dpi=100)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples).convert("L")
                img_np = np.array(img)
                result = run_easyocr(img_np)
                text += " ".join([txt for _, txt, _ in result]) + "\n"

                del img, img_np, result, pix
                gc.collect()
            except Exception as e:
                st.warning(f"OCR failed on page {i+1}: {e}")
        del doc
        gc.collect()
    except Exception as e:
        st.error(f"PDF OCR processing failed: {e}")
    return text.strip()


def docx_extract(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join(para.text for para in doc.paragraphs).strip()
        return text or ocr_docx_extract(file_path)
    except Exception as e:
        st.error(f"DOCX text extraction failed: {e}")
        return ocr_docx_extract(file_path)


def ocr_docx_extract(file_path):
    text_chunks = []
    try:
        with ZipFile(file_path) as docx_zip:
            for image_name in docx_zip.namelist():
                if image_name.startswith("word/media/"):
                    with docx_zip.open(image_name) as img_file:
                        try:
                            image = Image.open(BytesIO(img_file.read())).convert("L")
                            image.load()
                            image.thumbnail((1200, 1200), Image.LANCZOS)
                            img_np = np.array(image)

                            result = run_easyocr(img_np)
                            chunk = " ".join([txt for _, txt, _ in result])
                            text_chunks.append(chunk)

                            del image, img_np, result, chunk
                            gc.collect()
                        except Exception as e:
                            st.warning(f"OCR failed on image {image_name}: {e}")
    except Exception as e:
        st.error(f"DOCX OCR failed: {e}")
    return "\n".join(text_chunks).strip()


def txt_extract(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except Exception as e:
        st.error(f"TXT extraction failed: {e}")
        return ""


def clean_text(text):
    text = re.sub(r"\n+", "\n", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"[^\w\s.,;:?!'-]", '', text)
    return text.strip()


def extract(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    author = None
    creation_date = None

    if ext == ".pdf":
        try:
            with fitz.open(file_path) as doc:
                meta = doc.metadata
                author = meta.get("author") or meta.get("Author")
                creation_date = meta.get("creationDate") or meta.get("CreationDate")
                creation_date = decode_pdf_date(creation_date) if creation_date else None
        except Exception as e:
            st.warning(f"Metadata extraction from PDF failed: {e}")

        text = pdf_extract(file_path)
        if not text:
            text = ocr_pdf_extract(file_path)

    elif ext == ".docx":
        try:
            doc = Document(file_path)
            props = doc.core_properties
            author = props.author
            creation_date = str(props.created)
            creation_date = decode_pdf_date(creation_date) if creation_date else None
        except Exception as e:
            st.warning(f"Metadata extraction from DOCX failed: {e}")

        text = docx_extract(file_path)

    elif ext == ".txt":
        text = txt_extract(file_path)

    else:
        raise ValueError("Unsupported file format.")

    text = clean_text(text)

    return {
        "text": text,
        "author": author or "Unknown",
        "creation_date": creation_date or "Unknown"
    }

